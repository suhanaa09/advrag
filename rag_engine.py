"""
rag_engine.py  –  Enhanced RAG Engine
  • Web scraping    : requests + BeautifulSoup
  • Document upload : PDF (PyMuPDF/pdfplumber), DOCX (python-docx), TXT
  • Embeddings      : sentence-transformers (all-MiniLM-L6-v2, runs locally, free)
  • Vector store    : FAISS (in-memory)
  • LLM             : Groq API (llama / mixtral / gemma)
  • Live web search : Tavily Search API (FREE tier = 1000 searches/month)
  • Conversation memory : rolling window with summary compression
  • Multilingual    : Hindi, Hinglish, English auto-detect & respond
"""

from __future__ import annotations
import re
import time
import textwrap
import io
from typing import List, Dict, Any, Optional, Tuple
from urllib.parse import urljoin, urlparse

import requests
from bs4 import BeautifulSoup
import numpy as np
import faiss
from sentence_transformers import SentenceTransformer
from groq import Groq


# ─────────────────────────────────────────────────────────────────────────────
# Helpers
# ─────────────────────────────────────────────────────────────────────────────

def clean_text(text: str) -> str:
    # Preserve Devanagari (Hindi) characters along with ASCII
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def chunk_text(text: str, chunk_size: int = 500, overlap: int = 80) -> List[str]:
    words = text.split()
    chunks, i = [], 0
    while i < len(words):
        chunk = " ".join(words[i : i + chunk_size])
        if chunk:
            chunks.append(chunk)
        i += chunk_size - overlap
    return chunks


def detect_language(text: str) -> str:
    """Detect if text is Hindi, Hinglish, or English."""
    # Check for Devanagari script
    devanagari = re.findall(r'[\u0900-\u097F]', text)
    total_chars = len([c for c in text if c.isalpha()])
    
    if not total_chars:
        return "english"
    
    devanagari_ratio = len(devanagari) / max(total_chars, 1)
    
    if devanagari_ratio > 0.5:
        return "hindi"
    elif devanagari_ratio > 0.1:
        return "hinglish"
    
    # Check for common Hindi romanized words
    hinglish_markers = [
        r'\b(kya|hai|hain|nahi|nahin|aur|lekin|toh|yeh|woh|kaise|kyun|mujhe|tumhe|aap|main|hum)\b',
        r'\b(bata|batao|samjhao|chahiye|lagta|lagti|karo|karna|dena|lena)\b',
        r'\b(accha|theek|sahi|galat|pata|malum|समझ|bilkul|zaroor)\b'
    ]
    
    text_lower = text.lower()
    hinglish_count = sum(len(re.findall(p, text_lower)) for p in hinglish_markers)
    
    if hinglish_count >= 2:
        return "hinglish"
    
    return "english"


# ─────────────────────────────────────────────────────────────────────────────
# Document Parser
# ─────────────────────────────────────────────────────────────────────────────

class DocumentParser:
    """Parse PDF, DOCX, and TXT files into text."""

    @staticmethod
    def parse_pdf(file_bytes: bytes, filename: str) -> Tuple[str, str]:
        """Returns (text, method_used)"""
        # Try pdfplumber first (better text extraction)
        try:
            import pdfplumber
            with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                pages_text = []
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        pages_text.append(text)
            if pages_text:
                return "\n\n".join(pages_text), "pdfplumber"
        except ImportError:
            pass
        except Exception:
            pass

        # Fallback to PyMuPDF (fitz)
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(stream=file_bytes, filetype="pdf")
            pages_text = []
            for page in doc:
                pages_text.append(page.get_text())
            doc.close()
            if pages_text:
                return "\n\n".join(pages_text), "pymupdf"
        except ImportError:
            pass
        except Exception:
            pass

        raise ValueError(
            f"Could not parse '{filename}'. Install pdfplumber or PyMuPDF:\n"
            "pip install pdfplumber  OR  pip install pymupdf"
        )

    @staticmethod
    def parse_docx(file_bytes: bytes, filename: str) -> str:
        try:
            from docx import Document
            doc = Document(io.BytesIO(file_bytes))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            # Also extract table text
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        paragraphs.append(row_text)
            return "\n\n".join(paragraphs)
        except ImportError:
            raise ValueError(
                f"python-docx not installed. Run: pip install python-docx"
            )

    @staticmethod
    def parse_txt(file_bytes: bytes) -> str:
        for enc in ["utf-8", "utf-16", "latin-1", "cp1252"]:
            try:
                return file_bytes.decode(enc)
            except UnicodeDecodeError:
                continue
        return file_bytes.decode("utf-8", errors="replace")

    def parse(self, file_bytes: bytes, filename: str) -> str:
        """Auto-detect format and parse."""
        ext = filename.lower().rsplit(".", 1)[-1]
        
        if ext == "pdf":
            text, _ = self.parse_pdf(file_bytes, filename)
        elif ext in ("docx", "doc"):
            text = self.parse_docx(file_bytes, filename)
        elif ext == "txt":
            text = self.parse_txt(file_bytes)
        else:
            # Try as text
            text = self.parse_txt(file_bytes)
        
        return clean_text(text)


# ─────────────────────────────────────────────────────────────────────────────
# Conversation Memory
# ─────────────────────────────────────────────────────────────────────────────

class ConversationMemory:
    """Rolling conversation memory with context compression."""

    def __init__(self, max_turns: int = 8, summary_threshold: int = 6):
        self.messages: List[Dict[str, str]] = []  # only user/assistant roles
        self.summary: str = ""
        self.max_turns = max_turns
        self.summary_threshold = summary_threshold

    def add(self, role: str, content: str):
        """Add a user or assistant message."""
        assert role in ("user", "assistant"), f"Bad role: {role}"
        self.messages.append({"role": role, "content": content})

    def get_history_messages(self) -> List[Dict[str, str]]:
        """Return only user/assistant history for appending after system prompt."""
        # Keep last max_turns pairs (each pair = 2 messages)
        return self.messages[-(self.max_turns * 2):]

    def get_summary_text(self) -> str:
        """Return summary string to inject into the system prompt."""
        return self.summary

    def compress(self, groq_client, model: str):
        """Summarize old messages when history gets long."""
        if len(self.messages) <= self.summary_threshold * 2:
            return

        # Summarize everything except the most recent turns
        old_messages = self.messages[:-(self.max_turns * 2)]
        if not old_messages:
            return

        convo_text = "\n".join(
            f"{m['role'].upper()}: {m['content'][:300]}" for m in old_messages
        )
        summary_prompt = (
            f"Summarize this conversation briefly in 2-3 sentences, "
            f"preserving key facts, medical info asked, and user preferences:\n\n{convo_text}"
        )
        try:
            resp = groq_client.chat.completions.create(
                model=model,
                messages=[{"role": "user", "content": summary_prompt}],
                max_tokens=200,
                temperature=0.1,
            )
            new_summary = resp.choices[0].message.content.strip()
            # Append to existing summary if present
            self.summary = (
                f"{self.summary} | {new_summary}" if self.summary else new_summary
            )
            # Drop the old messages we summarized
            self.messages = self.messages[-(self.max_turns * 2):]
        except Exception:
            pass  # fail silently — don't break the main query

    def clear(self):
        self.messages.clear()
        self.summary = ""

    @property
    def turn_count(self) -> int:
        return len([m for m in self.messages if m["role"] == "user"])


# ─────────────────────────────────────────────────────────────────────────────
# Live Web Search  (Tavily — free tier)
# ─────────────────────────────────────────────────────────────────────────────

class LiveSearch:
    ENDPOINT = "https://api.tavily.com/search"

    def __init__(self, api_key: str):
        self.api_key = api_key

    def search(self, query: str, max_results: int = 5) -> List[Dict]:
        payload = {
            "api_key": self.api_key,
            "query": query,
            "search_depth": "basic",
            "max_results": max_results,
            "include_answer": True,
        }
        try:
            resp = requests.post(self.ENDPOINT, json=payload, timeout=15)
            resp.raise_for_status()
            data = resp.json()
            results = []
            if data.get("answer"):
                results.append({
                    "title": "Tavily Direct Answer",
                    "url": "tavily://answer",
                    "content": data["answer"],
                })
            for r in data.get("results", []):
                results.append({
                    "title": r.get("title", ""),
                    "url": r.get("url", ""),
                    "content": r.get("content", ""),
                })
            return results
        except Exception as e:
            return [{"title": "Search Error", "url": "", "content": str(e)}]


# ─────────────────────────────────────────────────────────────────────────────
# Web Scraper
# ─────────────────────────────────────────────────────────────────────────────

class WebScraper:
    HEADERS = {
        "User-Agent": (
            "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
            "AppleWebKit/537.36 (KHTML, like Gecko) "
            "Chrome/124.0 Safari/537.36"
        )
    }

    def scrape_url(self, url: str, depth: int = 1) -> Dict[str, Any]:
        visited: set[str] = set()
        all_pages: List[Dict] = []
        self._crawl(url, url, depth, visited, all_pages)
        return {"pages": all_pages, "total": len(all_pages)}

    def _crawl(self, base: str, url: str, depth: int, visited: set, pages: list):
        if url in visited or depth < 0:
            return
        visited.add(url)
        try:
            resp = requests.get(url, headers=self.HEADERS, timeout=15)
            resp.raise_for_status()
        except Exception:
            return

        soup = BeautifulSoup(resp.text, "html.parser")
        for tag in soup(["script", "style", "nav", "footer", "header", "aside", "form"]):
            tag.decompose()

        title = soup.title.string.strip() if soup.title and soup.title.string else url
        body = soup.get_text(separator=" ")
        text = clean_text(body)

        if len(text) > 100:
            pages.append({"url": url, "title": title, "text": text})

        if depth > 1:
            for a in soup.find_all("a", href=True):
                href = urljoin(base, a["href"])
                if urlparse(href).netloc == urlparse(base).netloc:
                    self._crawl(base, href, depth - 1, visited, pages)
                    time.sleep(0.3)


# ─────────────────────────────────────────────────────────────────────────────
# FAISS Vector Store
# ─────────────────────────────────────────────────────────────────────────────

class VectorStore:
    def __init__(self, dim: int = 384):
        self.dim = dim
        self.index = faiss.IndexFlatL2(dim)
        self.chunks: List[str] = []
        self.metadata: List[Dict] = []

    def add(self, embeddings: np.ndarray, chunks: List[str], sources: List[str]):
        self.index.add(embeddings.astype("float32"))
        self.chunks.extend(chunks)
        self.metadata.extend([{"source": s} for s in sources])

    def search(self, query_embedding: np.ndarray, top_k: int = 4):
        if self.index.ntotal == 0:
            return [], []
        q = query_embedding.astype("float32").reshape(1, -1)
        k = min(top_k, self.index.ntotal)
        distances, indices = self.index.search(q, k)
        results, srcs = [], []
        for i, idx in enumerate(indices[0]):
            if idx != -1:
                results.append(self.chunks[idx])
                srcs.append(self.metadata[idx]["source"])
        return results, srcs

    def clear(self):
        self.index.reset()
        self.chunks.clear()
        self.metadata.clear()

    @property
    def total(self):
        return self.index.ntotal


# ─────────────────────────────────────────────────────────────────────────────
# RAG Engine
# ─────────────────────────────────────────────────────────────────────────────

MULTILINGUAL_SYSTEM_BASE = """
You are a friendly, intelligent AI assistant — warm, natural, and human-like in conversation.

CONVERSATION RULES (follow strictly):
• If the user says hi / hello / hey / namaste or any greeting → just greet them back naturally and briefly. Do NOT ask what you can help with in a robotic way. Keep it casual: "Hey! 👋 What's up?" or "Hi there! How can I help?"
• For casual small talk (how are you, what's your name, etc.) → respond naturally like a friend, keep it short.
• For questions → give clear, helpful answers using the context provided.
• ALWAYS read the conversation history before replying — never ignore what was said before.
• If the user says "tell me more", "explain that", "go on" → continue from your last answer using history.
• Never repeat yourself. Never ignore prior context.

RESPONSE STYLE:
• Be warm and natural — never robotic, never stiff
• Use markdown: **bold** for key terms, bullet lists when genuinely helpful, code blocks for code
• Match the user's tone: casual → relaxed; technical → detailed
• Never start with "Certainly!", "Of course!", "Absolutely!" or "Great question!"
• Keep greetings SHORT — one line max

LANGUAGE RULES:
• Hindi (Devanagari script) → reply fully in Hindi
• Hinglish (Roman Hindi words mixed in) → reply in Hinglish naturally
• English → reply in English
• NEVER switch language unless the user does
"""


class RAGEngine:
    def __init__(
        self,
        groq_api_key: str,
        tavily_api_key: str = "",
        model: str = "llama-3.3-70b-versatile",
        top_k: int = 4,
        temperature: float = 0.4,
    ):
        self.groq_client = Groq(api_key=groq_api_key)
        self.model = model
        self.top_k = top_k
        self.temperature = temperature
        self.tavily_api_key = tavily_api_key
        self.live_search: Optional[LiveSearch] = (
            LiveSearch(tavily_api_key) if tavily_api_key else None
        )

        # Use ONNX backend — avoids loading PyTorch (saves ~800MB on Streamlit Cloud)
        # Falls back to default backend if ONNX runtime not available
        try:
            self.embedder = SentenceTransformer(
                "all-MiniLM-L6-v2",
                backend="onnx",
                model_kwargs={"file_name": "model.onnx"},
            )
        except Exception:
            self.embedder = SentenceTransformer("all-MiniLM-L6-v2")
        self.vector_store = VectorStore(dim=384)
        self.scraper = WebScraper()
        self.doc_parser = DocumentParser()
        self.memory = ConversationMemory(max_turns=10, summary_threshold=8)
        self._source_set: set[str] = set()

    def set_tavily_key(self, key: str):
        self.tavily_api_key = key
        self.live_search = LiveSearch(key)

    # ── Ingestion ─────────────────────────────────────────────────────────────

    def add_url(self, url: str, depth: int = 1) -> Dict[str, Any]:
        result = self.scraper.scrape_url(url, depth=depth)
        pages = result["pages"]
        if not pages:
            raise ValueError(f"No content extracted from {url}")

        total_chunks = 0
        for page in pages:
            chunks = chunk_text(page["text"])
            if not chunks:
                continue
            embeddings = self.embedder.encode(chunks, show_progress_bar=False)
            sources = [page["url"]] * len(chunks)
            self.vector_store.add(np.array(embeddings), chunks, sources)
            self._source_set.add(page["url"])
            total_chunks += len(chunks)

        return {"chunks": total_chunks, "pages": len(pages)}

    def add_text(self, text: str, label: str = "manual") -> Dict[str, Any]:
        text = clean_text(text)
        chunks = chunk_text(text)
        if not chunks:
            raise ValueError("No usable text found.")
        embeddings = self.embedder.encode(chunks, show_progress_bar=False)
        sources = [label] * len(chunks)
        self.vector_store.add(np.array(embeddings), chunks, sources)
        self._source_set.add(label)
        return {"chunks": len(chunks)}

    def add_document(self, file_bytes: bytes, filename: str) -> Dict[str, Any]:
        """Parse and index an uploaded document (PDF, DOCX, TXT)."""
        text = self.doc_parser.parse(file_bytes, filename)
        if len(text.strip()) < 50:
            raise ValueError(f"Could not extract meaningful text from '{filename}'")
        
        chunks = chunk_text(text)
        if not chunks:
            raise ValueError("No usable content found in document.")
        
        embeddings = self.embedder.encode(chunks, show_progress_bar=False)
        label = filename
        sources = [label] * len(chunks)
        self.vector_store.add(np.array(embeddings), chunks, sources)
        self._source_set.add(label)
        
        # Estimate pages for PDF
        page_count = text.count("\n\n") // 3 + 1
        
        return {
            "chunks": len(chunks),
            "chars": len(text),
            "filename": filename,
            "estimated_pages": page_count,
        }

    # ── Query pipeline ────────────────────────────────────────────────────────

    def query(self, question: str) -> Dict[str, Any]:
        # 1. Detect language
        lang = detect_language(question)

        # 2. Detect if this is a greeting / small talk (skip RAG retrieval)
        greeting_patterns = [
            r'^\s*(hi|hey|hello|hii+|helo|howdy|sup|yo)\s*[!.?]*\s*$',
            r'^\s*(namaste|namaskar|sat sri akal|jai hind)\s*[!.?]*\s*$',
            r'^\s*(how are you|how r u|kaisa hai|kya haal|kya hal|aap kaise|tum kaise)\s*[!.?]*\s*$',
            r'^\s*(good morning|good evening|good night|good afternoon)\s*[!.?]*\s*$',
            r'^\s*(thanks|thank you|shukriya|dhanyawad|thnx|thx)\s*[!.?]*\s*$',
            r'^\s*(bye|goodbye|alvida|cya|see you)\s*[!.?]*\s*$',
        ]
        is_greeting = any(re.search(p, question.strip(), re.IGNORECASE) for p in greeting_patterns)

        # 3. Enrich query with recent context for better retrieval on follow-ups
        history = self.memory.get_history_messages()
        if history and len(question.split()) < 8 and not is_greeting:
            last_user = next((m["content"] for m in reversed(history) if m["role"] == "user"), "")
            retrieval_query = f"{last_user} {question}" if last_user else question
        else:
            retrieval_query = question

        # 4. FAISS retrieval (skip for greetings)
        chunks, sources = [], []
        if not is_greeting:
            q_emb = self.embedder.encode([retrieval_query], show_progress_bar=False)[0]
            chunks, sources = self.vector_store.search(q_emb, top_k=self.top_k)

        used_web_search = False

        # 5. Build context block + system instruction
        summary_text = self.memory.get_summary_text()
        memory_note = (
            f"\n\nEARLIER CONVERSATION SUMMARY:\n{summary_text}"
            if summary_text else ""
        )

        if chunks:
            context = "\n\n---\n\n".join(
                f"[Source: {s}]\n{c}" for c, s in zip(chunks, sources)
            )
            system_prompt = (
                MULTILINGUAL_SYSTEM_BASE
                + memory_note
                + f"""

RETRIEVED DOCUMENT CONTEXT (answer primarily from this):
{context}

INSTRUCTIONS:
- Answer the user's question using the context above.
- If the question is a follow-up (e.g. "tell me more", "explain that", "what about..."), use BOTH the conversation history AND the context.
- If context covers it fully, use it. If only partially, add your own knowledge.
- Cite the source naturally where relevant.
- NEVER say "I don't have information" if the context contains relevant content.
"""
            )

        elif self.live_search:
            used_web_search = True
            web_results = self.live_search.search(question, max_results=5)
            context_parts = []
            for r in web_results:
                if r["content"]:
                    context_parts.append(
                        f"[Source: {r['url']}]\nTitle: {r['title']}\n{r['content']}"
                    )
                    sources.append(r["url"])
            context = "\n\n---\n\n".join(context_parts) if context_parts else "No results found."
            system_prompt = (
                MULTILINGUAL_SYSTEM_BASE
                + memory_note
                + f"""

LIVE WEB SEARCH RESULTS (fetched in real-time):
{context}

INSTRUCTIONS:
- Answer using the web results above.
- You have up-to-date information — do NOT say your knowledge is limited.
- Use conversation history for follow-up questions.
- Cite sources naturally.
"""
            )

        else:
            system_prompt = (
                MULTILINGUAL_SYSTEM_BASE
                + memory_note
                + """

No documents loaded and web search is off. Answer from your training knowledge.
Be helpful and natural. Don't blame limitations — just answer as best you can.
Use conversation history if the question is a follow-up.
"""
            )

        # 5. Build final messages list: [system] + [history] + [current user msg]
        messages = (
            [{"role": "system", "content": system_prompt}]
            + history
            + [{"role": "user", "content": question}]
        )

        # 6. Call Groq
        chat = self.groq_client.chat.completions.create(
            model=self.model,
            messages=messages,
            temperature=self.temperature,
            max_tokens=1536,
        )
        answer = chat.choices[0].message.content

        # 7. Save to memory
        self.memory.add("user", question)
        self.memory.add("assistant", answer)

        # 8. Compress if needed
        if self.memory.turn_count >= self.memory.summary_threshold:
            self.memory.compress(self.groq_client, self.model)

        unique_sources = [
            s for s in list(dict.fromkeys(sources))
            if s and s != "tavily://answer"
        ]

        return {
            "answer": answer,
            "sources": unique_sources,
            "chunks_used": len(chunks),
            "used_web_search": used_web_search,
            "language_detected": lang,
            "memory_turns": self.memory.turn_count,
        }

    # ── Utilities ─────────────────────────────────────────────────────────────

    def get_stats(self) -> Dict[str, int]:
        return {
            "chunks": self.vector_store.total,
            "sources": len(self._source_set),
            "memory_turns": self.memory.turn_count,
        }

    def clear_memory(self):
        """Clear conversation history only."""
        self.memory.clear()

    def clear(self):
        """Clear everything."""
        self.vector_store.clear()
        self._source_set.clear()
        self.memory.clear()
